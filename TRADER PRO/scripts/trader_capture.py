#!/usr/bin/env python3
"""
TRADER PRO — Captura Automática Multi-Timeframe
Metodología: ICT / SMC
Autor: Claude Code para Carlos Saavedra
"""

import asyncio
import json
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional

from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════
# CONFIGURACIÓN
# ═══════════════════════════════════════════

BASE_DIR       = Path("/Users/carlossaavedra/Documents/CARLOS SAAVEDRA/CLAUDE/TRADER PRO/Análisis")
CHROME_PROFILE = Path("/Users/carlossaavedra/Library/Application Support/Google/Chrome")
CHROME_BIN     = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
COOKIES_FILE   = Path(__file__).parent / "tv_cookies.json"

TIMEFRAMES = [
    {"label": "Semanal", "interval": "W",   "display": "1W"},
    {"label": "Diario",  "interval": "D",   "display": "1D"},
    {"label": "H4",      "interval": "240", "display": "4H"},
    {"label": "H1",      "interval": "60",  "display": "1H"},
    {"label": "15min",   "interval": "15",  "display": "15"},
    {"label": "5min",    "interval": "5",   "display": "5"},
]

INITIAL_WAIT    = 15   # segundos para la primera carga
CHART_LOAD_WAIT = 10   # segundos por cada cambio de timeframe
KEEP_PNGS       = True # Mantener capturas junto al Word


# ═══════════════════════════════════════════
# FUNCIONES DE SOPORTE
# ═══════════════════════════════════════════

def ask_asset() -> str:
    print("\n" + "═" * 55)
    print("  TRADER PRO — Captura Automática Multi-Timeframe")
    print("═" * 55)
    print("\nEjemplos: EURUSD  XAUUSD  US30  NQ1!  BTCUSDT\n")
    asset = input("  ¿Qué activo analizar? → ").strip().upper()
    if not asset:
        asset = "XAUUSD"
        print(f"  (Sin entrada — usando {asset} por defecto)")
    return asset


async def close_popups(page) -> None:
    selectors = [
        '[data-name="close-button"]',
        'button[aria-label="Close"]',
        '.js-dialog-close',
        '[class*="closeButton"]',
        'button:has-text("Aceptar")',
        'button:has-text("Accept")',
        'button:has-text("Got it")',
        'button:has-text("Dismiss")',
        'button:has-text("Maybe later")',
    ]
    for sel in selectors:
        try:
            btn = page.locator(sel).first
            if await btn.is_visible():
                await btn.click(timeout=1500)
                await asyncio.sleep(0.4)
        except Exception:
            pass
    try:
        await page.keyboard.press("Escape")
        await asyncio.sleep(0.3)
    except Exception:
        pass


async def wait_for_chart(page, timeout: int = 15) -> bool:
    chart_selectors = [
        'canvas',
        '[data-name="legend"]',
        '[class*="chart-toolbar"]',
        '[data-name="indicators-button"]',
        '[class*="paneControls"]',
    ]
    for sel in chart_selectors:
        try:
            el = page.locator(sel).first
            await el.wait_for(state="visible", timeout=timeout * 1000)
            return True
        except Exception:
            pass
    return False


async def ensure_logged_in(page) -> None:
    await asyncio.sleep(3)
    if await wait_for_chart(page, timeout=10):
        print("  ✅ TradingView cargado correctamente")
        return
    print("\n" + "─" * 55)
    print("  🔑  ACCION REQUERIDA: haz login en el navegador")
    print("  Luego presiona ENTER aquí.")
    print("─" * 55)
    input("  → ENTER cuando veas el gráfico → ")
    await asyncio.sleep(3)
    try:
        cookies = await page.context.cookies()
        tv_cookies = [c for c in cookies if "tradingview" in c.get("domain", "")]
        if tv_cookies:
            with open(COOKIES_FILE, "w") as f:
                json.dump(tv_cookies, f, indent=2)
            print("  💾 Sesión guardada para próximas veces")
    except Exception:
        pass


async def capture_chart(page, asset: str, tf: dict, session_dir: Path) -> Path:
    url = f"https://www.tradingview.com/chart/?symbol={asset}&interval={tf['interval']}"
    print(f"  📸 [{tf['display']}] {tf['label']} — navegando...")
    await page.goto(url, wait_until="domcontentloaded", timeout=30000)
    await asyncio.sleep(CHART_LOAD_WAIT)
    await close_popups(page)
    await wait_for_chart(page, timeout=12)
    await asyncio.sleep(2)
    img_path = session_dir / f"{asset}_{tf['label']}_{tf['display']}.png"
    await page.screenshot(path=str(img_path), full_page=False)
    kb = img_path.stat().st_size / 1024 if img_path.exists() else 0
    tag = "✅" if kb > 50 else "⚠️ "
    print(f"  {tag} Guardado → {img_path.name}  ({kb:.0f} KB)")
    return img_path


# ═══════════════════════════════════════════
# CAPTURA PLAYWRIGHT
# ═══════════════════════════════════════════

async def _capture_with_context(context, mode: str, asset: str, session_dir: Path) -> list:
    page = await context.new_page()
    if mode == "playwright_cookies" and COOKIES_FILE.exists():
        with open(COOKIES_FILE) as f:
            await context.add_cookies(json.load(f))
        print("  🍪 Sesión anterior cargada")
    print(f"\n  🌐 Cargando TradingView — {asset}...")
    await page.goto(
        f"https://www.tradingview.com/chart/?symbol={asset}&interval=D",
        wait_until="domcontentloaded", timeout=30000
    )
    await asyncio.sleep(INITIAL_WAIT)
    await close_popups(page)
    await ensure_logged_in(page)
    if mode == "playwright_cookies":
        try:
            cookies = await context.cookies()
            tv = [c for c in cookies if "tradingview" in c.get("domain", "")]
            with open(COOKIES_FILE, "w") as f:
                json.dump(tv, f, indent=2)
            print(f"  💾 Sesión guardada → {COOKIES_FILE.name}")
        except Exception:
            pass
    print(f"\n  📷 Capturando {len(TIMEFRAMES)} timeframes:")
    captured = []
    for tf in TIMEFRAMES:
        img_path = await capture_chart(page, asset, tf, session_dir)
        captured.append({"tf": tf, "path": img_path})
    await context.close()
    return captured


async def run_capture(asset: str) -> tuple:
    fecha = datetime.now().strftime("%Y-%m-%d_%H-%M")
    session_dir = BASE_DIR / f"{asset}_{fecha}"
    session_dir.mkdir(parents=True, exist_ok=True)
    captured = []
    async with async_playwright() as p:
        mode_a_ok = False
        try:
            print(f"\n🚀 Modo A — Chrome con perfil de usuario...")
            ctx_a = await p.chromium.launch_persistent_context(
                user_data_dir=str(CHROME_PROFILE),
                executable_path=CHROME_BIN,
                headless=False,
                viewport={"width": 1920, "height": 1080},
                args=["--start-maximized", "--disable-blink-features=AutomationControlled",
                      "--disable-infobars", "--profile-directory=Default"]
            )
            captured = await _capture_with_context(ctx_a, "chrome_profile", asset, session_dir)
            mode_a_ok = True
        except Exception as e:
            print(f"  ⚠️  Modo A falló ({type(e).__name__}). Usando Modo B...")
        if not mode_a_ok:
            browser = await p.chromium.launch(
                headless=False,
                args=["--start-maximized", "--disable-blink-features=AutomationControlled"]
            )
            ctx_b = await browser.new_context(viewport={"width": 1920, "height": 1080})
            captured = await _capture_with_context(ctx_b, "playwright_cookies", asset, session_dir)
    return captured, session_dir, fecha


# ═══════════════════════════════════════════
# WORD — HELPERS (formato exacto XAUUSD ref)
# ═══════════════════════════════════════════

def _hex_to_rgb(color_hex: str) -> RGBColor:
    b = bytes.fromhex(color_hex)
    return RGBColor(b[0], b[1], b[2])


def _p(doc, text: str = "", bold: bool = False, size_pt: float = 11,
       color_hex: str = "", align=None,
       space_before_pt: float = -1, italic: bool = False) -> object:
    """Crea un párrafo con formato exacto — sin Heading styles."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    if space_before_pt >= 0:
        para.paragraph_format.space_before = Pt(space_before_pt)
    if align is not None:
        para.alignment = align
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        if color_hex:
            run.font.color.rgb = _hex_to_rgb(color_hex)
    return para


def _shade_cell(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _cell(cell, text: str, bold: bool = False, size_pt: float = 10,
          color_hex: Optional[str] = None, bg_hex: Optional[str] = None) -> None:
    """Formatea una celda de tabla."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = _hex_to_rgb(color_hex)
    if bg_hex:
        _shade_cell(cell, bg_hex)


def _set_col_widths(table, widths_inches: list) -> None:
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_inches):
                cell.width = Inches(widths_inches[i])


# ═══════════════════════════════════════════
# WORD — GENERACIÓN (formato idéntico al ref)
# ═══════════════════════════════════════════

def generate_word(captured: list, session_dir: Path, asset: str, fecha: str,
                  analysis: Optional[dict] = None) -> Path:
    """
    Genera el Word con formato exactamente igual al archivo de referencia
    XAUUSD_2026-03-15_23-41_ANALISIS.docx.
    """
    doc = Document()
    an = analysis or {}

    # Márgenes — 0.75" en todos los lados (igual que el ref)
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)

    # Eliminar el párrafo vacío inicial que crea Document() por defecto
    # (preservar sectPr — propiedades de sección)
    sectPr_tag = qn('w:sectPr')
    for elem in list(doc.element.body):
        if elem.tag != sectPr_tag:
            doc.element.body.remove(elem)

    now_str = datetime.now().strftime("%d/%m/%Y  %H:%M")
    asset_display = an.get("asset_display", asset)

    # ── PORTADA ──────────────────────────────────────────────────────────────
    _p(doc, f"ANÁLISIS DE MERCADO — {asset_display}",
       bold=True, size_pt=22, color_hex="1F3864",
       align=WD_ALIGN_PARAGRAPH.CENTER)

    _p(doc, f"{now_str}  |  Bogotá",
       size_pt=13, color_hex="646464",
       align=WD_ALIGN_PARAGRAPH.CENTER)

    _p(doc, "─" * 60,
       color_hex="1F3864",
       align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── SECCIÓN 1: SESGO ─────────────────────────────────────────────────────
    _p(doc, "SECCIÓN 1: ¿QUÉ ESTÁ PASANDO HOY EN EL MERCADO?",
       bold=True, size_pt=14, color_hex="1F3864",
       align=WD_ALIGN_PARAGRAPH.CENTER)

    _p(doc,
       "Esta sección resume en palabras simples lo que está haciendo el mercado en "
       "diferentes escalas de tiempo — desde la tendencia de meses hasta lo que pasa ahora mismo.",
       space_before_pt=6)

    # Tabla sesgo (6 filas x 3 columnas)
    sesgo_filas = an.get("sesgo_filas", [
        ("Tendencia de meses (Gráfico Semanal)", "[ SUBE ↑ / BAJA ↓ ]",
         "La dirección principal del mercado en los últimos meses."),
        ("Últimos días (Gráfico Diario)",        "[ SUBE ↑ / BAJA ↓ ]",
         "Lo que ha pasado esta semana."),
        ("Últimas horas (Gráfico de 4 horas)",   "[ SUBE ↑ / BAJA ↓ ]",
         "El movimiento de las últimas horas."),
        ("Ahora mismo (Gráfico de 1 hora)",      "[ SUBE ↑ / BAJA ↓ ]",
         "Lo que está haciendo el precio en este momento."),
        ("RECOMENDACIÓN DEL DÍA",               "[ VENDER / COMPRAR ]",
         "Lo que el análisis sugiere hacer hoy."),
    ])
    tbl_sesgo = doc.add_table(rows=len(sesgo_filas) + 1, cols=3)
    tbl_sesgo.style = "Table Grid"
    _set_col_widths(tbl_sesgo, [2.00, 1.30, 3.20])
    for j, h in enumerate(["Escala de Tiempo", "¿Sube o Baja?", "¿Qué significa?"]):
        _cell(tbl_sesgo.rows[0].cells[j], h, bold=True,
              color_hex="FFFFFF", bg_hex="1F3864")
    for i, (escala, dir_, que) in enumerate(sesgo_filas):
        bg = "EBF0FA" if i % 2 == 0 else "FFFFFF"
        _cell(tbl_sesgo.rows[i+1].cells[0], escala, bold=True, bg_hex=bg)
        _cell(tbl_sesgo.rows[i+1].cells[1], dir_,   bg_hex=bg)
        _cell(tbl_sesgo.rows[i+1].cells[2], que,    bg_hex=bg)

    _p(doc)  # separador

    # ── SECCIÓN 2: GRÁFICOS ───────────────────────────────────────────────────
    _p(doc, "SECCIÓN 2: LOS GRÁFICOS — LO QUE SE VE EN CADA UNO",
       bold=True, size_pt=14, color_hex="1F3864",
       space_before_pt=14)

    tf_titles = {
        "Semanal": "GRÁFICO SEMANAL (1W)",
        "Diario":  "GRÁFICO DIARIO (1D)",
        "H4":      "GRÁFICO DE 4 HORAS (4H)",
        "H1":      "GRÁFICO DE 1 HORA (1H)",
        "15min":   "GRÁFICO DE 15 MINUTOS (15)",
        "5min":    "GRÁFICO DE 5 MINUTOS (5)",
    }
    tf_analisis = an.get("tf_analisis", {})

    for item in captured:
        tf    = item["tf"]
        path  = item["path"]
        titulo: str = tf_titles.get(tf["label"]) or tf["label"]
        texto: str  = tf_analisis.get(tf["label"]) or "Lo que se ve: pendiente — compartir imagen con Claude Code."

        _p(doc, titulo, bold=True, size_pt=12, color_hex="1F3864", space_before_pt=12)

        # Imagen centrada
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para.paragraph_format.space_before = Pt(6)
        pic_para.paragraph_format.space_after  = Pt(0)
        pic_para.paragraph_format.line_spacing = 1.0
        if path.exists():
            run_pic = pic_para.add_run()
            run_pic.add_picture(str(path), width=Inches(6.5))
        else:
            pic_para.add_run(f"[Imagen no encontrada: {path.name}]")

        # Análisis del TF
        ana_para = doc.add_paragraph()
        ana_para.paragraph_format.space_before = Pt(4)
        ana_para.paragraph_format.space_after  = Pt(0)
        ana_para.paragraph_format.line_spacing = 1.0
        ana_para.add_run(texto).font.size = Pt(10.5)

    _p(doc)  # separador

    # ── SECCIÓN 3: OPERACIÓN ──────────────────────────────────────────────────
    _p(doc, "SECCIÓN 3: ¿QUÉ HACER? — LA OPERACIÓN RECOMENDADA",
       bold=True, size_pt=14, color_hex="1F3864", space_before_pt=14)

    _p(doc,
       "Aquí están los números exactos. Lee cada fila con calma. "
       "La operación solo es válida si se cumplen las condiciones de la Sección 5.",
       space_before_pt=6)

    _p(doc)  # línea en blanco antes de tabla

    op_filas = an.get("operacion_filas", [
        ("¿COMPRAR o VENDER?",
         "VENDER — apostamos a que el precio seguirá bajando"),
        ("¿EN QUÉ PRECIO ENTRAS?",
         "[ precio exacto ] → Solo entra cuando el precio llegue exactamente a ese número"),
        ("¿CUÁNDO SALIR SI PIERDES? (Stop Loss)",
         "[ precio exacto ] → Si el precio llega aquí, CIERRA la operación. Pérdida pequeña y controlada."),
        ("¿CUÁNDO TOMAR LA PRIMERA GANANCIA? (Objetivo 1)",
         "[ precio exacto ] → Cuando el precio llegue aquí, cierra la mitad. Ya aseguraste ganancia."),
        ("¿CUÁNDO TOMAR TODA LA GANANCIA? (Objetivo 2)",
         "[ precio exacto ] → Cuando llegue aquí, cierra el resto. Es la ganancia máxima del análisis."),
        ("¿CUÁNTO PUEDES GANAR?",
         "Objetivo 1: [ X ] veces lo arriesgado  /  Objetivo 2: [ X ] veces lo arriesgado"),
        ("¿CUÁNDO CANCELAR TODO?",
         "[ precio exacto ] → Si el precio cierra una vela de 4 horas POR ENCIMA de este nivel: el análisis ya no es válido."),
        ("¿A QUÉ HORA ENTRAR?",
         "Horario de mayor actividad en Bogotá: Mañana 9am–12pm  /  Tarde 2pm–6pm"),
    ])
    tbl_op = doc.add_table(rows=len(op_filas), cols=2)
    tbl_op.style = "Table Grid"
    _set_col_widths(tbl_op, [2.50, 4.00])
    for i, (campo, val) in enumerate(op_filas):
        _cell(tbl_op.rows[i].cells[0], campo, bold=True, bg_hex="D9E1F2")
        _cell(tbl_op.rows[i].cells[1], val)

    _p(doc)  # separador

    # ── SECCIÓN 4: EXPLICACIÓN ────────────────────────────────────────────────
    _p(doc, "SECCIÓN 4: EXPLICACIÓN EN PALABRAS SIMPLES",
       bold=True, size_pt=14, color_hex="1F3864", space_before_pt=14)

    _p(doc, "¿POR QUÉ ESTA OPERACIÓN?",
       bold=True, size_pt=12, color_hex="1F3864", space_before_pt=8)

    _p(doc, "Aquí la explicación completa sin ningún tecnicismo:", space_before_pt=2)

    explicaciones = an.get("explicaciones", [
        ("¿QUÉ PASÓ?",
         "Pendiente de análisis — compartir imágenes con Claude Code."),
        ("¿POR QUÉ ESTA ZONA?",
         "Pendiente de análisis — compartir imágenes con Claude Code."),
        ("¿QUÉ SE ESPERA?",
         "Pendiente de análisis — compartir imágenes con Claude Code."),
        ("¿QUÉ SEÑAL ESPERAR ANTES DE ENTRAR?",
         "Pendiente de análisis — compartir imágenes con Claude Code."),
    ])
    for label, texto in explicaciones:
        bul = doc.add_paragraph()
        bul.paragraph_format.space_before = Pt(4)
        bul.paragraph_format.space_after  = Pt(0)
        bul.paragraph_format.line_spacing = 1.0
        bul.add_run("• ")
        bul.add_run(f"{label}: ").bold = True
        bul.add_run(texto)

    _p(doc)
    _p(doc)

    # ── SECCIÓN 5: LISTA DE VERIFICACIÓN ─────────────────────────────────────
    _p(doc, "SECCIÓN 5: LISTA DE VERIFICACIÓN — ANTES DE ENTRAR",
       bold=True, size_pt=14, color_hex="1F3864", space_before_pt=14)

    _p(doc,
       "Marca cada punto antes de hacer clic en 'vender'. "
       "Si no puedes marcar al menos 5 de los 8 puntos, NO entres a la operación.",
       space_before_pt=6)

    checklist = an.get("checklist", [
        "☐  ¿La tendencia de 4 horas y 1 hora van en la misma dirección (las dos bajando)?",
        "☐  ¿El precio llegó a la zona de entrada exacta definida en la Sección 3?",
        "☐  ¿El precio hizo un pico rápido hacia arriba y luego volvió a bajar? → Trampa para compradores",
        "☐  ¿Estás en el horario de alta actividad? → 9am–12pm o 2pm–6pm hora Bogotá",
        "☐  ¿Apareció una vela roja grande en el gráfico de 15 minutos en esa zona?",
        "☐  ¿El precio está subiendo hacia una zona 'cara'? → Zona premium para vender",
        "☐  ¿El volumen es alto en la zona de venta? (si tienes ese indicador) → Volumen alto = movimiento real",
        "☐  ¿El flujo de órdenes (CVD) está bajando? (si tienes ese indicador) → CVD bajando = vendedores dominan",
    ])
    for item in checklist:
        chk = doc.add_paragraph()
        chk.paragraph_format.space_before = Pt(3)
        chk.paragraph_format.space_after  = Pt(0)
        chk.paragraph_format.line_spacing = 1.0
        chk.add_run(item).font.size = Pt(10.5)

    _p(doc)

    # ── SECCIÓN 6: GESTIÓN DEL DINERO ─────────────────────────────────────────
    _p(doc, "SECCIÓN 6: ¿CUÁNTO ARRIESGAR? — GESTIÓN DEL DINERO",
       bold=True, size_pt=14, color_hex="1F3864", space_before_pt=14)

    _p(doc)  # línea en blanco
    _p(doc)

    riesgo_filas = an.get("riesgo_filas", [
        ("REGLA DE ORO",
         "Nunca arriesgues más del 1% de tu capital en una sola operación. "
         "Con $10,000 → máximo pierdes $100."),
        ("¿CÓMO CALCULAR LOS LOTES?",
         "Fórmula: Capital × 1% ÷ distancia al Stop Loss = tamaño de posición"),
        ("EJEMPLO PARA ESTA OPERACIÓN",
         "[ Completar con datos exactos del activo ]"),
        ("SI PIERDES 3 VECES SEGUIDAS",
         "Para de operar ese día. Revisa el análisis. Las pérdidas seguidas "
         "indican que el mercado cambió o hay algo que no estás viendo."),
        ("PÉRDIDA MÁXIMA DEL DÍA",
         "Si en un día pierdes el 3% de tu capital, CIERRA TODO y descansa. "
         "El mercado sigue mañana."),
    ])
    tbl_riesgo = doc.add_table(rows=len(riesgo_filas), cols=2)
    tbl_riesgo.style = "Table Grid"
    _set_col_widths(tbl_riesgo, [2.50, 4.00])
    for i, (campo, val) in enumerate(riesgo_filas):
        _cell(tbl_riesgo.rows[i].cells[0], campo, bold=True, bg_hex="D9E1F2")
        _cell(tbl_riesgo.rows[i].cells[1], val)

    _p(doc)
    _p(doc)

    # ── PIE DE PÁGINA ─────────────────────────────────────────────────────────
    _p(doc,
       f"TRADER PRO — Análisis {asset_display} | {now_str} | Bogotá",
       size_pt=9, color_hex="828282",
       align=WD_ALIGN_PARAGRAPH.CENTER,
       space_before_pt=20)

    # Guardar
    docx_path = session_dir / f"{asset}_{fecha}_ANALISIS.docx"
    doc.save(str(docx_path))
    return docx_path


# ═══════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════

async def main() -> None:
    if len(sys.argv) > 1:
        asset = sys.argv[1].strip().upper()
        print(f"\n{'═'*55}")
        print(f"  TRADER PRO — Activo: {asset}  (argumento CLI)")
        print(f"{'═'*55}")
    else:
        asset = ask_asset()

    print(f"\n  Activo    : {asset}")
    print(f"  Timeframes: {', '.join(tf['display'] for tf in TIMEFRAMES)}")

    # 1. Capturar screenshots
    captured, session_dir, fecha = await run_capture(asset)

    print(f"\n  📁 Imágenes en:\n     {session_dir}")

    # 2. Verificar capturas
    ok = sum(1 for c in captured if c["path"].exists() and c["path"].stat().st_size > 50000)
    print(f"\n  🔍 Verificación: {ok}/{len(captured)} imágenes OK (>50KB)")
    for c in captured:
        p = c["path"]
        kb = p.stat().st_size / 1024 if p.exists() else 0
        tag = "✅" if kb > 50 else "⚠️ "
        print(f"     {tag} {p.name}  ({kb:.0f} KB)")

    # 3. Cargar analysis.json si existe
    analysis_json = session_dir / "analysis.json"
    if not analysis_json.exists():
        analysis_json = Path(__file__).parent / "analysis.json"
    analysis = None
    if analysis_json.exists():
        try:
            with open(analysis_json) as f:
                analysis = json.load(f)
            print(f"\n  📊 Análisis cargado desde analysis.json")
        except Exception:
            pass

    # 4. Generar Word
    print("\n  📝 Generando documento Word...")
    docx_path = generate_word(captured, session_dir, asset, fecha, analysis)
    print(f"  ✅ Word guardado → {docx_path.name}  ({docx_path.stat().st_size/1024:.0f} KB)")

    # 5. PNGs se mantienen
    if KEEP_PNGS:
        print(f"\n  📌 Capturas PNG mantenidas en carpeta para revisión.")

    # 6. Resumen
    print(f"\n{'═'*55}")
    print("  ANÁLISIS COMPLETADO")
    print(f"{'═'*55}")
    print(f"  Activo  : {asset}")
    print(f"  Sesión  : {fecha}")
    print(f"  Archivo : {docx_path.name}")
    print(f"  Carpeta : {session_dir}")
    print(f"{'═'*55}\n")
    print("  SIGUIENTE PASO: comparte las imágenes PNG con Claude Code")
    print("  para recibir el análisis ICT/SMC completo.")

    # 7. Abrir carpeta y Word
    os.system(f'open "{session_dir}"')
    os.system(f'open "{docx_path}"')


if __name__ == "__main__":
    asyncio.run(main())
